import streamlit as st
import os
import tempfile
import shutil
from pathlib import Path
from typing import List, Tuple, Dict
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
import docx2txt
import json
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain_community.vectorstores import FAISS
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains import create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.documents import Document
from dotenv import load_dotenv
from datetime import datetime
import asyncio
import aiofiles
import logging
import hashlib
from concurrent.futures import ThreadPoolExecutor
import threading

# Load environment variables
load_dotenv()

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Custom exceptions
class APIKeyError(Exception):
    pass

class VectorStoreError(Exception):
    pass

# Thread pool executor for synchronous API calls
executor = ThreadPoolExecutor(max_workers=4)

# Event loop for async tasks
loop = asyncio.new_event_loop()
threading.Thread(target=loop.run_forever, daemon=True).start()

def run_async(coroutine):
    future = asyncio.run_coroutine_threadsafe(coroutine, loop)
    return future.result()

# File vector stores and metadata
file_vectorstores = {}
file_metadata = {}

# Page configuration
st.set_page_config(
    page_title="Document Q&A Chatbot",
    page_icon="🤖",
    layout="wide"
)

# Initialize session state
if 'vectorstore' not in st.session_state:
    st.session_state.vectorstore = None
if 'qa_chain' not in st.session_state:
    st.session_state.qa_chain = None
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'documents_processed' not in st.session_state:
    st.session_state.documents_processed = False
if 'api_key_valid' not in st.session_state:
    st.session_state.api_key_valid = False
if 'file_list' not in st.session_state:
    st.session_state.file_list = []

async def validate_api_key_with_test():
    google_key = os.getenv('GOOGLE_API_KEY')
    if not google_key:
        raise APIKeyError("Google API key is required. Please set GOOGLE_API_KEY in your .env file.")
    
    def test_embedding():
        embeddings = GoogleGenerativeAIEmbeddings(
            model="models/text-embedding-004",
            google_api_key=google_key
        )
        embeddings.embed_query("test")
        return google_key

    try:
        google_key = await asyncio.get_event_loop().run_in_executor(executor, test_embedding)
        logger.info("API key validated successfully")
        return google_key
    except Exception as e:
        error_msg = str(e).lower()
        if any(keyword in error_msg for keyword in ['api key', 'authentication', 'unauthorized', 'forbidden', 'invalid']):
            raise APIKeyError(f"Invalid or expired Google API key. Error: {str(e)}")
        elif 'quota' in error_msg or 'limit' in error_msg:
            raise APIKeyError(f"API quota exceeded. Error: {str(e)}")
        else:
            raise APIKeyError(f"API key validation failed: {str(e)}")

async def extract_text_from_file(file_path: str, file_ext: str) -> str:
    try:
        text = ""
        if file_ext == ".pdf":
            async with aiofiles.open(file_path, "rb") as f:
                content = await f.read()
            reader = PdfReader(file_path)
            text = "".join([page.extract_text() or '' for page in reader.pages])
        elif file_ext == ".docx":
            try:
                async with aiofiles.open(file_path, "rb") as f:
                    content = await f.read()
                text = docx2txt.process(file_path)
            except:
                doc = DocxDocument(file_path)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
        elif file_ext == ".doc":
            try:
                import textract
                text = textract.process(file_path).decode("utf-8")
            except ImportError:
                st.warning("⚠️ textract not available for .doc files. Install with: pip install textract")
                logger.warning("textract not installed for .doc file processing")
                return ""
            except Exception as e:
                st.error(f"Error processing .doc file: {str(e)}")
                logger.error(f"Failed to process .doc file: {str(e)}")
                return ""
        elif file_ext == ".txt":
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            for encoding in encodings:
                try:
                    async with aiofiles.open(file_path, "r", encoding=encoding) as f:
                        text = await f.read()
                    break
                except UnicodeDecodeError:
                    continue
        else:
            st.warning(f"Unsupported file format: {file_ext}")
            logger.warning(f"Unsupported file format: {file_ext}")
            return ""
        text = text.strip()
        if not text:
            st.warning(f"No text content extracted from {file_path}")
            logger.warning(f"No text content extracted from {file_path}")
            return ""
        logger.info(f"Successfully extracted {len(text)} characters from {file_path}")
        return text
    except Exception as e:
        st.error(f"Failed to extract text from {file_path}: {str(e)}")
        logger.error(f"Failed to extract text from {file_path}: {str(e)}")
        return ""

async def create_file_vectorstore(document: Document, google_api_key: str, file_id: str):
    try:
        def create_embeddings():
            return GoogleGenerativeAIEmbeddings(
                model="models/text-embedding-004",
                google_api_key=google_api_key
            )
        
        embeddings = await asyncio.get_event_loop().run_in_executor(executor, create_embeddings)
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=1500,
            chunk_overlap=10,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        chunks = []
        doc_chunks = splitter.split_text(document.page_content)
        for i, chunk in enumerate(doc_chunks):
            if chunk.strip():
                chunks.append(Document(
                    page_content=chunk,
                    metadata={
                        **document.metadata,
                        "chunk_id": i,
                        "total_chunks": len(doc_chunks),
                        "file_id": file_id
                    }
                ))
        if not chunks:
            st.warning(f"No chunks created for {document.metadata.get('source_file')}")
            logger.warning(f"No chunks created for {document.metadata.get('source_file')}")
            return None
        def create_vector_store():
            vectorstore = FAISS.from_documents(chunks, embeddings)
            vectorstore.save_local(f"faiss_index_{file_id}")
            return vectorstore
        vectorstore = await asyncio.get_event_loop().run_in_executor(executor, create_vector_store)
        logger.info(f"Created {len(chunks)} chunks for file {document.metadata['source_file']}")
        return vectorstore
    except Exception as e:
        st.error(f"Failed to create vector store for file: {str(e)}")
        logger.error(f"Failed to create vector store for file: {str(e)}")
        return None

async def create_combined_vectorstore(google_api_key: str):
    global file_vectorstores
    try:
        if not file_vectorstores:
            return None
        def create_embeddings():
            return GoogleGenerativeAIEmbeddings(
                model="models/text-embedding-004",
                google_api_key=google_api_key
            )
        embeddings = await asyncio.get_event_loop().run_in_executor(executor, create_embeddings)
        file_ids = list(file_vectorstores.keys())
        combined_vectorstore = file_vectorstores[file_ids[0]]
        for file_id in file_ids[1:]:
            combined_vectorstore.merge_from(file_vectorstores[file_id])
        def save_combined():
            combined_vectorstore.save_local("faiss_index_combined")
        await asyncio.get_event_loop().run_in_executor(executor, save_combined)
        logger.info(f"Created combined vector store from {len(file_vectorstores)} files")
        return combined_vectorstore
    except Exception as e:
        st.error(f"Failed to create combined vector store: {str(e)}")
        logger.error(f"Failed to create combined vector store: {str(e)}")
        return None

async def load_file_vectorstores(google_api_key: str) -> Dict:
    global file_vectorstores, file_metadata
    try:
        def create_embeddings():
            return GoogleGenerativeAIEmbeddings(
                model="models/text-embedding-004",
                google_api_key=google_api_key
            )
        embeddings = await asyncio.get_event_loop().run_in_executor(executor, create_embeddings)
        if os.path.exists("file_metadata.json"):
            async with aiofiles.open("file_metadata.json", "r") as f:
                file_metadata = json.loads(await f.read())
        loaded_stores = {}
        for file_id in file_metadata.keys():
            file_index_dir = f"faiss_index_{file_id}"
            if os.path.exists(file_index_dir):
                try:
                    def load_vector_store():
                        return FAISS.load_local(
                            file_index_dir,
                            embeddings,
                            allow_dangerous_deserialization=True
                        )
                    vectorstore = await asyncio.get_event_loop().run_in_executor(executor, load_vector_store)
                    loaded_stores[file_id] = vectorstore
                    logger.info(f"Loaded vector store for file: {file_metadata[file_id]['filename']}")
                except Exception as e:
                    st.error(f"Failed to load vector store for {file_id}: {str(e)}")
                    logger.error(f"Failed to load vector store for {file_id}: {str(e)}")
        file_vectorstores = loaded_stores
        return loaded_stores
    except Exception as e:
        st.error(f"Failed to load file vector stores: {str(e)}")
        logger.error(f"Failed to load file vector stores: {str(e)}")
        return {}

async def load_combined_vectorstore(google_api_key: str):
    try:
        def create_embeddings():
            return GoogleGenerativeAIEmbeddings(
                model="models/text-embedding-004",
                google_api_key=google_api_key
            )
        embeddings = await asyncio.get_event_loop().run_in_executor(executor, create_embeddings)
        if not os.path.exists("faiss_index_combined"):
            return None
        def load_vector_store():
            return FAISS.load_local(
                "faiss_index_combined",
                embeddings,
                allow_dangerous_deserialization=True
            )
        vectorstore = await asyncio.get_event_loop().run_in_executor(executor, load_vector_store)
        logger.info("Combined vector store loaded successfully")
        return vectorstore
    except Exception as e:
        st.error(f"Failed to load combined vector store: {str(e)}")
        logger.error(f"Failed to load combined vector store: {str(e)}")
        return None

async def build_qa_chain(vectorstore, google_api_key: str):
    try:
        if not vectorstore:
            raise VectorStoreError("Cannot build QA chain: vector store is None")
        def create_llm():
            return ChatGoogleGenerativeAI(
                model="gemini-1.5-flash",
                google_api_key=google_api_key,
                temperature=0.2,
                max_output_tokens=800
            )
        llm = await asyncio.get_event_loop().run_in_executor(executor, create_llm)
        system_prompt = """
        You are a helpful document assistant. Use the provided context to answer questions accurately.
        Guidelines:
        - Answer based solely on the provided context
        - Be concise but complete
        - If information is not in the context, clearly state that
        - Maintain a professional tone
        Context: {context}
        """
        prompt = ChatPromptTemplate.from_messages([
            ("system", system_prompt),
            ("human", "{input}")
        ])
        retriever = vectorstore.as_retriever(
            search_type="similarity_score_threshold",
            search_kwargs={
                "k": 8,
                "score_threshold": 0.3
            }
        )
        document_chain = create_stuff_documents_chain(llm, prompt)
        qa_chain = create_retrieval_chain(retriever, document_chain)
        logger.info("QA chain built successfully")
        return qa_chain
    except Exception as e:
        st.error(f"Failed to build QA chain: {str(e)}")
        logger.error(f"Failed to build QA chain: {str(e)}")
        raise VectorStoreError(f"Failed to build QA chain: {str(e)}")

async def process_uploaded_files(uploaded_files, google_api_key: str, current_user):
    global file_vectorstores, file_metadata
    temp_dir = tempfile.mkdtemp()
    processed_files = []
    try:
        logger.info(f"Processing {len(uploaded_files)} files separately...")
        for uploaded_file in uploaded_files:
            try:
                file_name = uploaded_file.name
                file_ext = Path(file_name).suffix.lower()
                file_path = os.path.join(temp_dir, file_name)
                async with aiofiles.open(file_path, "wb") as f:
                    await f.write(uploaded_file.getvalue())
                text = await extract_text_from_file(file_path, file_ext)
                if text and len(text.strip()) > 10:
                    document = Document(
                        page_content=text,
                        metadata={
                            "source_file": file_name,
                            "file_type": file_ext,
                            "file_size": len(uploaded_file.getvalue())
                        }
                    )
                    file_id = get_file_id(file_name)
                    vectorstore = await create_file_vectorstore(document, google_api_key, file_id)
                    if vectorstore:
                        file_vectorstores[file_id] = vectorstore
                        file_metadata[file_id] = {
                            "filename": file_name,
                            "file_type": file_ext,
                            "file_size": len(uploaded_file.getvalue()),
                            "text_length": len(text),
                            "uploaded_by": current_user.id,
                            "user_email": current_user.email,
                            "accessible_to_all": False,
                            "processed_at": datetime.utcnow().isoformat()
                        }
                        processed_files.append(file_name)
                        logger.info(f"Created separate vector store for: {file_name}")
                    else:
                        st.warning(f"Failed to create vector store for: {file_name}")
                        logger.warning(f"Failed to create vector store for: {file_name}")
                else:
                    st.warning(f"Skipped {file_name} - insufficient content")
                    logger.warning(f"Skipped {file_name} - insufficient content")
            except Exception as e:
                st.error(f"Failed to process {file_name}: {str(e)}")
                logger.error(f"Failed to process {file_name}: {str(e)}")
                continue
        if not processed_files:
            raise VectorStoreError("No valid documents were processed")
        combined_vectorstore = await create_combined_vectorstore(google_api_key)
        qa_chain = await build_qa_chain(combined_vectorstore, google_api_key)
        logger.info(f"Successfully processed {len(processed_files)} files with vector stores")
        return combined_vectorstore, qa_chain
    except Exception as e:
        st.error(f"Failed to process files: {str(e)}")
        logger.error(f"Failed to process files: {str(e)}")
        return None, None
    finally:
        try:
            shutil.rmtree(temp_dir, ignore_errors=True)
        except Exception as e:
            logger.error(f"Error cleaning up temp directory: {str(e)}")

def get_file_id(filename: str) -> str:
    return hashlib.md5(filename.encode()).hexdigest()[:8]

async def answer_question(qa_chain, vectorstore, question: str, google_api_key: str):
    try:
        if not question.strip():
            return {"answer": "Please provide a valid question."}
        logger.info(f"Processing question: {question[:100]}...")
        def invoke_qa_chain():
            return qa_chain.invoke({"input": question})
        response = await asyncio.get_event_loop().run_in_executor(executor, invoke_qa_chain)
        answer = response.get("answer", "I couldn't find relevant information in the documents.")
        return {"answer": answer}
    except Exception as e:
        st.error(f"Failed to process question: {str(e)}")
        logger.error(f"Failed to process question: {str(e)}")
        return {"answer": f"Error processing question: {str(e)}"}

def save_metadata():
    global file_metadata
    try:
        with open("file_metadata.json", "w") as f:
            json.dump(file_metadata, f, indent=2)
        logger.info("File metadata saved successfully")
    except Exception as e:
        st.error(f"Failed to save metadata: {str(e)}")
        logger.error(f"Failed to save metadata: {str(e)}")

def cleanup_file_data(file_id: str):
    global file_vectorstores, file_metadata
    try:
        if file_id in file_vectorstores:
            del file_vectorstores[file_id]
        if file_id in file_metadata:
            del file_metadata[file_id]
        file_index_dir = f"faiss_index_{file_id}"
        if os.path.exists(file_index_dir):
            shutil.rmtree(file_index_dir, ignore_errors=True)
        save_metadata()
        logger.info(f"Cleaned up data for file: {file_id}")
    except Exception as e:
        st.error(f"Failed to cleanup file data: {str(e)}")
        logger.error(f"Failed to cleanup file data: {str(e)}")

def validate_api_key():
    try:
        google_key = run_async(validate_api_key_with_test())
        st.session_state.api_key_valid = True
        return google_key
    except APIKeyError as e:
        st.error(f"🔑 API Key Error: {str(e)}")
        st.info("💡 Please set GOOGLE_API_KEY in your .env file and ensure it's valid.")
        st.session_state.api_key_valid = False
        return None
    except Exception as e:
        st.error(f"❌ Unexpected error during API key validation: {str(e)}")
        st.session_state.api_key_valid = False
        return None

def main():
    st.title("🤖 Document Q&A Chatbot")
    st.write("Upload PDF, DOCX, DOC, or TXT files and ask questions about their content.")

    with st.sidebar:
        st.subheader("🔑 API Key Status")
        google_api_key = validate_api_key()
        if google_api_key:
            st.success("✅ API Key Valid")
            run_async(load_file_vectorstores(google_api_key))
            st.session_state.vectorstore = run_async(load_combined_vectorstore(google_api_key))
            if st.session_state.vectorstore:
                st.session_state.qa_chain = run_async(build_qa_chain(st.session_state.vectorstore, google_api_key))
                st.session_state.documents_processed = True
        else:
            st.stop()

        st.divider()
        st.subheader("📁 Upload Files")
        uploaded_files = st.file_uploader(
            "Choose files",
            type=['pdf', 'docx', 'doc', 'txt'],
            accept_multiple_files=True,
            help="Upload PDF, DOCX, DOC, or TXT files to analyze"
        )
        if uploaded_files and st.button("🚀 Process Documents", type="primary"):
            try:
                with st.spinner("Processing documents..."):
                    current_user = type('User', (), {'id': 'user123', 'email': 'user@example.com'})()
                    vectorstore, qa_chain = run_async(process_uploaded_files(uploaded_files, google_api_key, current_user))
                    if vectorstore and qa_chain:
                        st.session_state.vectorstore = vectorstore
                        st.session_state.qa_chain = qa_chain
                        st.session_state.documents_processed = True
                        st.session_state.file_list = [
                            {
                                "file_id": file_id,
                                "filename": meta["filename"],
                                "file_type": meta["file_type"],
                                "file_size": meta["file_size"],
                                "text_length": meta.get("text_length"),
                                "uploaded_by": meta.get("uploaded_by"),
                                "user_email": meta.get("user_email"),
                                "accessible_to_all": meta.get("accessible_to_all", False)
                            }
                            for file_id, meta in file_metadata.items()
                        ]
                        st.success(f"✅ Successfully processed {len(uploaded_files)} files!")
                        st.rerun()
            except VectorStoreError as e:
                st.error(f"📚 Document Processing Error: {str(e)}")
            except Exception as e:
                st.error(f"❌ Unexpected error during document processing: {str(e)}")

        st.divider()
        st.subheader("📊 Status")
        if st.session_state.documents_processed:
            st.success("✅ Documents processed and ready")
            st.write("Processed Files:")
            for file in st.session_state.file_list:
                with st.expander(f"{file['filename']}"):
                    st.write(f"File Type: {file['file_type']}")
                    st.write(f"File Size: {file['file_size']} bytes")
                    st.write(f"Text Length: {file['text_length']} characters")
                    st.write(f"Uploaded By: {file['user_email']}")
        else:
            st.info("📄 Upload documents to start")

        if st.button("🗑️ Clear Session & Data"):
            try:
                for file_id in list(file_vectorstores.keys()):
                    cleanup_file_data(file_id)
                if os.path.exists("faiss_index_combined"):
                    shutil.rmtree("faiss_index_combined", ignore_errors=True)
                st.session_state.clear()
                st.success("🔄 Session and data cleared.")
                st.rerun()
            except Exception as e:
                st.error(f"❌ Error clearing data: {str(e)}")

    if not st.session_state.documents_processed:
        st.info("👈 Please upload and process documents using the sidebar to start chatting.")
    else:
        st.subheader("💬 Ask Questions")
        col1, col2 = st.columns([3, 1])
        with col1:
            user_question = st.text_area(
                "Enter your question:",
                placeholder="What would you like to know about the uploaded documents?",
                height=100,
                key="user_input"
            )
            if st.button("Ask", type="primary") and user_question:
                if st.session_state.qa_chain and st.session_state.vectorstore:
                    try:
                        with st.spinner("🤔 Analyzing documents..."):
                            response = run_async(answer_question(
                                st.session_state.qa_chain,
                                st.session_state.vectorstore,
                                user_question,
                                google_api_key
                            ))
                            st.subheader("💡 Answer")
                            answer_text = response.get('answer', 'No answer generated.')
                            if any(phrase in answer_text.lower() for phrase in ['not specified', 'not mentioned', 'not found', 'not provided', 'not available', 'i cannot find']):
                                st.warning("⚠️ Information may be incomplete or missing from the documents.")
                            elif "error" in answer_text.lower():
                                st.error("❌ An error occurred while processing your question.")
                            else:
                                st.success("✅ Answer found in documents.")
                            st.write(answer_text)
                            st.session_state.chat_history.append({
                                "question": user_question,
                                "answer": answer_text
                            })
                    except Exception as e:
                        st.error(f"❌ Error processing question: {str(e)}")
                else:
                    st.warning("⚠️ Please upload and process documents first.")
        
        with col2:
            if st.session_state.chat_history:
                st.subheader("📝 Chat History")
                for i, chat in enumerate(reversed(st.session_state.chat_history)):
                    with st.expander(f"Q: {chat['question'][:60]}{'...' if len(chat['question']) > 60 else ''}", expanded=(i == 0)):
                        st.markdown(f"**Question:** {chat['question']}")
                        st.markdown(f"**Answer:** {chat['answer']}")

if __name__ == "__main__":
    main()